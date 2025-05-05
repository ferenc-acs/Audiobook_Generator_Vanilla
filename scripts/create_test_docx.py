#
#    Audiobook Generator Vanilla: A Python application that converts text documents (DOCX, EPUB, Markdown) into audiobooks using OpenAI's text-to-speech API.
#    Copyright (C) 2025 Ferenc Acs <pass.schist2954@eagereverest.com>
#
#    This program is free software: you can redistribute it and/or modify
#    it under the terms of the GNU Affero General Public License as
#    published by the Free Software Foundation, either version 3 of the
#    License, or (at your option) any later version.
#
#    This program is distributed in the hope that it will be useful,
#    but WITHOUT ANY WARRANTY; without even the implied warranty of
#    MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
#    GNU Affero General Public License for more details.
#
#    You should have received a copy of the GNU Affero General Public License
#    along with this program.  If not, see <https://www.gnu.org/licenses/>.
#

from docx import Document
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE
import os
from pathlib import Path

def create_test_document_with_chapters(output_path="test_with_chapters.docx"):
    """Create a test DOCX file with multiple chapters using heading styles."""
    doc = Document()
    
    # Configure heading styles
    styles = doc.styles
    
    # Add title
    doc.add_heading('Sample Audiobook with Chapters', 0)
    doc.add_paragraph('This document demonstrates chapter segmentation for audiobook generation.')
    
    # Chapter 1
    doc.add_heading('Chapter 1: Introduction', 1)
    doc.add_paragraph('This is the first chapter of our sample audiobook. ')
    doc.add_paragraph('Chapter detection works by identifying heading styles in DOCX files. ')
    doc.add_paragraph('The audiobook generator should process this chapter separately and add an indicator at the end.')
    
    # Chapter 2
    doc.add_heading('Chapter 2: The Journey Begins', 1)
    doc.add_paragraph('In this chapter, our protagonist sets off on an adventure. ')
    doc.add_paragraph('The sun was rising over the distant mountains as they packed their belongings. ')
    doc.add_paragraph('"Today is going to be a great day," they thought to themselves.')
    
    # Chapter 3 with subsections
    doc.add_heading('Chapter 3: Discoveries', 1)
    doc.add_paragraph('This chapter contains multiple sections to test nested heading detection.')
    
    doc.add_heading('The Forest', 2)
    doc.add_paragraph('The forest was dense and mysterious. Sunlight filtered through the canopy in dappled patterns.')
    
    doc.add_heading('The Cave', 2)
    doc.add_paragraph('The cave entrance was hidden behind a curtain of vines. Inside, it was cool and dark.')
    
    # Chapter 4
    doc.add_heading('Chapter 4: Conclusion', 1)
    doc.add_paragraph('This is the final chapter of our test document. ')
    doc.add_paragraph('It demonstrates how the audiobook generator handles multiple chapters. ')
    doc.add_paragraph('Each chapter should be processed separately with an end-of-chapter indicator added.')
    
    # Save the document
    output_dir = Path(output_path).parent
    if not output_dir.exists():
        output_dir.mkdir(parents=True, exist_ok=True)
        
    doc.save(output_path)
    return output_path

def create_simple_test_document(output_path="test.docx"):
    """Create a simple test DOCX file without explicit chapter structure."""
    doc = Document()
    doc.add_paragraph('Sample audiobook text for integration testing')
    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Create both test documents
    simple_doc_path = create_simple_test_document()
    chapters_doc_path = create_test_document_with_chapters()
    
    print(f"Simple test document created: {simple_doc_path}")
    print(f"Chapter test document created: {chapters_doc_path}")
    print("\nRun the audiobook generator with these files to test chapter segmentation:")
    print(f"python -m src.main {chapters_doc_path}")