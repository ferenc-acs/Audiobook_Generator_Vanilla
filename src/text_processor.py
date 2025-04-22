from typing import Optional, List, Dict, Tuple
from docx import Document
from ebooklib import epub
import logging
import re
from bs4 import BeautifulSoup

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path: str) -> str:
    try:
        doc = Document(file_path)
        return '\n'.join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error(f"DOCX extraction error: {str(e)}")
        raise

def extract_chapters_from_docx(file_path: str) -> List[Dict[str, str]]:
    """Extract chapters from a DOCX file based on heading styles."""
    try:
        doc = Document(file_path)
        chapters = []
        current_chapter = {"title": "Untitled Chapter", "content": []}
        
        for para in doc.paragraphs:
            # Check if this paragraph is a heading (potential chapter marker)
            # Safely check if style exists before accessing its name
            if para.style and para.style.name.startswith('Heading') and para.text.strip():
                # If we have content in the current chapter, save it
                if current_chapter["content"]:
                    current_chapter["content"] = '\n'.join(current_chapter["content"])
                    chapters.append(current_chapter)
                
                # Start a new chapter
                current_chapter = {"title": para.text.strip(), "content": []}
            else:
                # Add content to the current chapter
                if para.text.strip():
                    current_chapter["content"].append(para.text)
        
        # Add the last chapter if it has content
        if current_chapter["content"]:
            current_chapter["content"] = '\n'.join(current_chapter["content"])
            chapters.append(current_chapter)
            
        # If no chapters were detected, treat the entire document as one chapter
        if not chapters:
            chapters.append({
                "title": "Complete Document",
                "content": '\n'.join([para.text for para in doc.paragraphs])
            })
            
        return chapters
    except Exception as e:
        logger.error(f"DOCX chapter extraction error: {str(e)}")
        raise

def extract_text_from_epub(file_path: str) -> str:
    """Extracts all text content from an EPUB file, using BeautifulSoup for robustness."""
    try:
        book = epub.read_epub(file_path)
        all_text_content = []
        # Iterate through items in spine order for better concatenation
        spine_ids = {item_id for item_id, _ in book.spine}
        items_to_process = []
        processed_ids = set()

        # Prioritize spine items
        for item_id, _ in book.spine:
            item = book.get_item_with_id(item_id)
            if item and item.media_type and item.media_type.lower() in ['application/xhtml+xml', 'text/html']:
                items_to_process.append(item)
                processed_ids.add(item.id)

        # Add non-spine items (like title pages, etc.) that might have been missed
        for item in book.get_items():
            if item.id not in processed_ids and item.media_type and item.media_type.lower() in ['application/xhtml+xml', 'text/html']:
                items_to_process.append(item)
                logger.debug(f"Adding non-spine item to text extraction: {item.get_name()}")

        for item in items_to_process:
            try:
                content_bytes = item.get_content()
                soup = BeautifulSoup(content_bytes, 'html.parser')
                # Extract text without initial stripping
                text_content = soup.get_text('\n') 
                stripped_content = text_content.strip()
                if stripped_content:
                    all_text_content.append(stripped_content)
            except Exception as item_e:
                logger.warning(f"Could not process item {item.get_name()} in EPUB for text extraction: {item_e}")
                
        return '\n\n'.join(all_text_content)
    except Exception as e:
        logger.error(f"EPUB extraction error: {str(e)}")
        raise

def extract_chapters_from_epub(file_path: str) -> List[Dict[str, str]]:
    """Extract chapters from an EPUB file."""
    try:
        book = epub.read_epub(file_path)
        chapters = []
        toc = []
        
        # Extract table of contents manually since get_toc method isn't consistently available
        try:
            # Look for the NCX file which contains the table of contents
            ncx_item = None
            for item in book.get_items():
                if item.get_name().endswith('.ncx'):
                    ncx_item = item
                    break
            
            if ncx_item:
                from lxml import etree
                ns = {'ncx': 'http://www.daisy.org/z3986/2005/ncx/'}
                tree = etree.fromstring(ncx_item.get_content())
                
                # Extract TOC entries from the NCX file
                nav_points = tree.xpath('//ncx:navPoint', namespaces=ns)
                toc = []
                
                for nav_point in nav_points:
                    # Get the title
                    title_elem = nav_point.find('.//ncx:text', namespaces=ns)
                    title = title_elem.text if title_elem is not None else 'Untitled'
                    
                    # Get the content source
                    content_elem = nav_point.find('.//ncx:content', namespaces=ns)
                    if content_elem is not None and 'src' in content_elem.attrib:
                        href = content_elem.attrib['src']
                        # Split off any fragment identifier
                        href = href.split('#')[0]
                        toc.append((title, href, []))
                
                logger.info(f"Extracted {len(toc)} entries from EPUB NCX table of contents")
            else:
                # If NCX not found, try to find the navigation document (EPUB3)
                nav_item = None
                for item in book.get_items():
                    if item.get_properties() and 'nav' in item.get_properties():
                        nav_item = item
                        break
                
                if nav_item:
                    soup = BeautifulSoup(nav_item.get_content(), 'html.parser')
                    toc = []
                    nav_element = soup.find('nav', attrs={'epub:type': 'toc'}) or soup.find('nav')
                    
                    if nav_element:
                        for link in nav_element.find_all('a'):
                            if 'href' in link.attrs:
                                href = link.attrs['href']
                                # Split off any fragment identifier
                                href = href.split('#')[0]
                                toc.append((link.get_text().strip(), href, []))
                    
                    logger.info(f"Extracted {len(toc)} entries from EPUB3 navigation document")
                else:
                    logger.info("Could not find TOC in NCX or navigation document")
                    toc = []
        except Exception as e:
            logger.warning(f"Failed to extract EPUB table of contents: {str(e)}")
            toc = []
        
        item_map = {}
        for item in book.get_items():
            if item.media_type and item.media_type.lower() in ['application/xhtml+xml', 'text/html']:
                item_map[item.get_name()] = item

        # Prioritize TOC for chapter structure
        if toc:
            logger.info(f"Using EPUB Table of Contents ({len(toc)} entries).")
            processed_items = set()
            for title, href, _ in toc:
                # Split href in case it includes fragment identifiers (#section)
                href_base = href.split('#')[0]
                
                if href_base in item_map:
                    item = item_map[href_base]
                    if href_base in processed_items:
                        logger.debug(f"Skipping duplicate TOC entry processing for: {href_base}")
                        continue # Avoid processing the same item multiple times if TOC is structured oddly
                    
                    # Decode content and clean HTML
                    content_bytes = item.get_content()
                    soup = BeautifulSoup(content_bytes, 'html.parser')
                    text_content = soup.get_text('\n', strip=True)
                    
                    # Check for common non-chapter titles like 'Contents', 'Title Page', 'Copyright'
                    lower_title = title.lower()
                    if any(keyword in lower_title for keyword in ['contents', 'title page', 'copyright', 'introduction', 'preface', 'dedication']):
                        logger.info(f"Skipping non-chapter TOC entry: {title} ({href_base})")
                        continue

                    # Check if content seems substantial enough to be a chapter
                    if len(text_content) > 100: # Arbitrary threshold to filter out very short/empty pages
                        chapters.append({
                            "title": title.strip(),
                            "content": text_content
                        })
                        processed_items.add(href_base)
                    else:
                        logger.debug(f"Skipping potentially empty/short TOC entry: {title} ({href_base}) - Length: {len(text_content)}")
                else:
                    logger.warning(f"TOC entry href not found in EPUB items: {href} (Base: {href_base})")

        # Fallback: Process HTML items in spine order if TOC was missing or yielded no chapters
        if not chapters:
            logger.warning("TOC empty or unusable. Falling back to spine item processing.")
            # Use book.spine for reading order
            spine_items = book.spine
            if not spine_items:
                 logger.warning("EPUB spine is empty. Cannot determine reading order for fallback.")
            else:
                logger.info(f"Processing {len(spine_items)} items from EPUB spine.")
                # Create a map of ID -> item for quick lookup
                id_map = {item.id: item for item in book.get_items() if item.media_type and item.media_type.lower() in ['application/xhtml+xml', 'text/html']}

                for i, (item_id, _) in enumerate(spine_items):
                    if item_id in id_map:
                        item = id_map[item_id]
                        item_name = item.get_name() # Get name for logging/fallback title
                        content_bytes = item.get_content()
                        soup = BeautifulSoup(content_bytes, 'html.parser')
                        # Get text without stripping whitespace initially
                        text_content = soup.get_text('\n') 

                        # Try to find a heading for the title, otherwise use item name or chapter number
                        title_elem = soup.find(['h1', 'h2', 'h3'])
                        title = f"Chapter {i+1}" # Default fallback title
                        if title_elem and title_elem.get_text().strip():
                            potential_title = title_elem.get_text().strip()
                            # Use heading as title if it's reasonably short
                            if len(potential_title) < 100:
                               title = potential_title
                            else:
                               logger.debug(f"Using default title; Found heading seems too long: {potential_title}")
                        elif item_name:
                           # Use item name if useful, clean it up
                           name_part = item_name.split('/')[-1].split('.')[0].replace('_', ' ').title()
                           if len(name_part) > 3 and len(name_part) < 50: # Avoid generic names like 'page1'
                               title = name_part

                        logger.debug(f"Spine fallback processing item: {item_name} (ID: {item_id}) -> Title: '{title}', Raw text length: {len(text_content)}")

                        # Add chapter if content exists (removed length check)
                        stripped_text = text_content.strip()
                        if stripped_text:
                           chapters.append({
                               "title": title,
                               "content": stripped_text # Use stripped text
                           })
                        else:
                           logger.debug(f"Skipping spine item (empty after strip): {title} (ID: {item_id})")
                    else:
                        logger.warning(f"Item ID '{item_id}' from spine not found in book items map.")
        
        # Final fallback: If no chapters were identified, treat the entire book text as one chapter.
        if not chapters:
            logger.warning("No chapters identified through TOC or spine fallback. Attempting to treat entire EPUB as one chapter.")
            try:
                all_text = extract_text_from_epub(file_path) # Re-extract all text cleanly
                logger.info(f"Final fallback: Extracted total text length: {len(all_text)}")
                stripped_all_text = all_text.strip()
                if stripped_all_text:
                    logger.info("Final fallback: Adding 'Complete Book' chapter.")
                    chapters.append({
                        "title": "Complete Book",
                        "content": stripped_all_text
                    })
                else:
                    logger.warning("Final fallback: Extracted text was empty after stripping. No chapters added.")
            except Exception as final_fallback_e:
                logger.error(f"Final fallback failed during text extraction: {final_fallback_e}")
        
        # Return the extracted chapters
        return chapters
    except Exception as e:
        logger.error(f"EPUB chapter extraction error: {str(e)}")
        raise

def process_text_file(file_path: str) -> Optional[str]:
    """Legacy method to extract all text from a file without chapter segmentation."""
    if file_path.endswith('.docx'):
        return extract_text_from_docx(file_path)
    if file_path.endswith('.epub'):
        return extract_text_from_epub(file_path)
    if file_path.lower().endswith(('.md', '.markdown')):
        return extract_text_from_markdown(file_path)
    return None

def extract_chapters(file_path: str) -> List[Dict[str, str]]:
    """Extract chapters from a document, segmenting by chapter boundaries."""
    if file_path.endswith('.docx'):
        return extract_chapters_from_docx(file_path)
    if file_path.endswith('.epub'):
        return extract_chapters_from_epub(file_path)
    if file_path.lower().endswith(('.md', '.markdown')):
        return extract_chapters_from_markdown(file_path)
    return []

def extract_text_from_markdown(file_path: str) -> str:
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Markdown extraction error: {str(e)}")
        raise

def extract_chapters_from_markdown(file_path: str) -> List[Dict[str, str]]:
    """Extract chapters from a Markdown file based on headings."""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            content = file.read()
            
        # Look for heading markers (# for h1, ## for h2, etc.)
        # We'll consider h1 or h2 as chapter headings
        chapter_pattern = re.compile(r'^(#{1,2})\s+(.+)$', re.MULTILINE)
        chapter_matches = list(chapter_pattern.finditer(content))
        
        chapters = []
        
        # If we found chapter headings
        if chapter_matches:
            for i, match in enumerate(chapter_matches):
                chapter_title = match.group(2).strip()
                start_pos = match.end()
                
                # Find the end of this chapter (start of next chapter or end of file)
                if i < len(chapter_matches) - 1:
                    end_pos = chapter_matches[i + 1].start()
                else:
                    end_pos = len(content)
                
                chapter_content = content[start_pos:end_pos].strip()
                chapters.append({
                    "title": chapter_title,
                    "content": chapter_content
                })
        
        # If no chapters were found, look for other patterns like "Chapter X" text
        if not chapters:
            chapter_text_pattern = re.compile(r'^(Chapter\s+\d+|Part\s+\d+|Book\s+\d+)(.*)$', re.MULTILINE | re.IGNORECASE)
            chapter_matches = list(chapter_text_pattern.finditer(content))
            
            if chapter_matches:
                for i, match in enumerate(chapter_matches):
                    chapter_title = match.group(0).strip()
                    start_pos = match.end()
                    
                    # Find the end of this chapter
                    if i < len(chapter_matches) - 1:
                        end_pos = chapter_matches[i + 1].start()
                    else:
                        end_pos = len(content)
                    
                    chapter_content = content[start_pos:end_pos].strip()
                    chapters.append({
                        "title": chapter_title,
                        "content": chapter_content
                    })
        
        # If still no chapters, treat the entire document as one chapter
        if not chapters:
            chapters.append({
                "title": "Complete Document",
                "content": content
            })
            
        return chapters
    except Exception as e:
        logger.error(f"Markdown chapter extraction error: {str(e)}")
        raise

def check_supported_format(file_path: str) -> bool:
    return file_path.lower().endswith(('.docx', '.epub', '.md', '.markdown'))